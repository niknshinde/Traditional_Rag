# Document Loaders
# ================
# This file handles reading different document types (PDF, DOCX, TXT)

from pathlib import Path

# ============================================
# PYTHON CONCEPT: pathlib.Path
# ============================================
# Path is the modern way to handle file paths in Python.
# It works on Windows, Mac, and Linux without you worrying about / vs \
#
# Examples:
#   path = Path("documents/file.pdf")
#   path.suffix  →  ".pdf"
#   path.stem    →  "file"
#   path.exists() → True/False


def load_pdf(file_path: Path) -> str:
    """
    Load text content from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as a string
    """
    # ============================================
    # PYTHON CONCEPT: Import inside function
    # ============================================
    # We import here (not at top) because:
    # 1. Not every user has PyPDF2 installed
    # 2. If they only use TXT files, no error occurs
    # This is called "lazy importing"
    
    try:
        from pypdf import PdfReader
    except ImportError:
        # Fallback to older package name
        from PyPDF2 import PdfReader
    
    reader = PdfReader(file_path)
    
    # ============================================
    # PYTHON CONCEPT: List Comprehension
    # ============================================
    # This is a compact way to build a list.
    # 
    # Long form:
    #   texts = []
    #   for page in reader.pages:
    #       texts.append(page.extract_text())
    #
    # Short form (list comprehension):
    texts = [page.extract_text() for page in reader.pages]
    
    # Join all pages with newlines
    return "\n".join(texts)


def load_docx(file_path: Path) -> str:
    """
    Load text content from a Word document.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text as a string
    """
    from docx import Document
    
    doc = Document(file_path)
    
    # Extract text from all paragraphs
    # doc.paragraphs is a list of Paragraph objects
    # Each paragraph has a .text attribute
    texts = [paragraph.text for paragraph in doc.paragraphs]
    
    return "\n".join(texts)


def load_txt(file_path: Path) -> str:
    """
    Load text content from a plain text file.
    
    Args:
        file_path: Path to the TXT file
        
    Returns:
        File contents as a string
    """
    # ============================================
    # PYTHON CONCEPT: Handling Multiple Encodings
    # ============================================
    # Different programs save files in different encodings:
    # - Most modern tools: UTF-8
    # - Windows PowerShell: UTF-16 (with BOM)
    # - Old Windows: Latin-1
    #
    # We try multiple encodings until one works!
    
    encodings_to_try = ['utf-8', 'utf-16', 'utf-16-le', 'latin-1']
    
    for encoding in encodings_to_try:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue  # Try next encoding
    
    # If all failed, raise an error
    raise ValueError(f"Could not decode file {file_path} with any known encoding")


def load_document(file_path: str) -> str:
    """
    Load a document based on its file extension.
    
    This is the MAIN function you'll use. It automatically
    detects the file type and calls the right loader.
    
    Args:
        file_path: Path to the document (string or Path)
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file type is not supported
    """
    # Convert string to Path object (if it isn't already)
    path = Path(file_path)
    
    # ============================================
    # PYTHON CONCEPT: Match statement (Python 3.10+)
    # ============================================
    # This is like switch/case in other languages.
    # .suffix gives us the file extension (e.g., ".pdf")
    # .lower() converts to lowercase for case-insensitive matching
    
    match path.suffix.lower():
        case ".pdf":
            return load_pdf(path)
        case ".docx":
            return load_docx(path)
        case ".txt":
            return load_txt(path)
        case _:
            # _ is the "catch-all" pattern
            raise ValueError(f"Unsupported file type: {path.suffix}")
